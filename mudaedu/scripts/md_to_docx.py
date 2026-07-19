#!/usr/bin/env python3
"""Converte Markdown simples (relatórios) para DOCX."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def add_formatted_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_done = False
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                if code_lang:
                    cap = doc.add_paragraph(f"[Diagrama: {code_lang}]")
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(9)
                code_lang = ""
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# ") and not title_done:
            h = doc.add_heading(stripped[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_done = True
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for col, header in enumerate(headers):
                cell = table.rows[0].cells[col]
                cell.text = ""
                add_formatted_runs(cell.paragraphs[0], header)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for r, row in enumerate(rows):
                for c, value in enumerate(row[: len(headers)]):
                    cell = table.rows[r + 1].cells[c]
                    cell.text = ""
                    add_formatted_runs(cell.paragraphs[0], value)
            doc.add_paragraph()
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("Se quiser, posso gerar"):
            break

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


if __name__ == "__main__":
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("docs/relatorio-infraestrutura-usuarios.md")
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(".docx")
    md_to_docx(src, dst)
    print(f"Gerado: {dst}")
